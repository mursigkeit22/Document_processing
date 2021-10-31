from docx import *
"""
https://github.com/python-openxml/python-docx/pull/716/commits/c09da22afaebaed2f0a3139de6ba46c8824f179e

docx/__init__.py
31 + PartFactory.part_type_for[CT.WML_DOCUMENT_MACRO_ENABLED_MAIN] = DocumentPart
docx/api.py
   - if document_part.content_type != CT.WML_DOCUMENT_MAIN:
26 + if (document_part.content_type != CT.WML_DOCUMENT_MAIN) and (document_part.content_type != CT.WML_DOCUMENT_MACRO_ENABLED_MAIN):

docx/opc/constants.py
283 +    WML_DOCUMENT_MACRO_ENABLED_MAIN = (
284 +        'application/vnd.ms-word.document.macroEnabled.main+xml'
285 +    )
"""

document = Document('./work_docs/doc_with_macros.docm')

for t in document.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.highlight_color is None:
                        run.font.hidden = True

for p in document.paragraphs:
    for run in p.runs:

        if run.font.highlight_color is None:
            run.font.hidden = True

document.save('./work_docs/done/doc_with_macros.docm')