from docx import *

document = Document('./text/full.docx')

for t in document.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.highlight_color is None:
                        run.font.hidden = True

for p in document.paragraphs:
    for run in p.runs:

        if run.font.highlight_color is None:
            run.font.hidden = True

document.save('./text/result_docx.docx')
