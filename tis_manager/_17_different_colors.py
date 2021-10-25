from docx import *
from docx.enum.text import WD_COLOR_INDEX
from typing import List

# colors_list = ["YELLOW", "BLACK", "BRIGHT_GREEN", "TEAL"]

document = Document('./text/different_colors.docx')


def return_pydocx_color(colors_list: List[str]):
    wd_colors_list = [getattr(WD_COLOR_INDEX, color) for color in colors_list]
    return wd_colors_list


for t in document.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    print(run.font.highlight_color)
                    # if run.font.highlight_color in return_pydocx_color(colors_list):
                    #     run.font.hidden = True

for p in document.paragraphs:
    for run in p.runs:
        if run.font.highlight_color:
            print(run.font.highlight_color)
            # print(run.font.highlight_color)
            # run.font.hidden = True

# document.save('./text/done/different_colors.docx')
