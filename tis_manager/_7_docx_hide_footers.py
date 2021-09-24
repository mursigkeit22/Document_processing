"""
https://python-docx.readthedocs.io/en/latest/api/section.html#sections-objects
"""
from docx import *

"""
нет доступа к боковым колонтитулам
"""

document = Document('./text/full.docx')


def iterate_footer_header(footer_or_header):
    for t in footer_or_header.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.highlight_color is None:
                            run.font.hidden = True
    for p in footer_or_header.paragraphs:
        for run in p.runs:
            if run.font.highlight_color is None:
                run.font.hidden = True


for section in document.sections:

    footers_and_headers = (section.footer, section.even_page_footer, section.first_page_footer,
                           section.header, section.even_page_header, section.first_page_header)
    for footer in footers_and_headers:
        iterate_footer_header(footer)

document.save('./text/result_footers.docx')
