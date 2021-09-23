"""
https://python-docx.readthedocs.io/en/latest/api/section.html#sections-objects
"""
from docx import *

document = Document('./text/full.docx')

for section in document.sections:
    footer = section.footer
    for t in footer.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print(p.text)
                    for run in p.runs:
                        if run.font.highlight_color is None:
                            run.font.hidden = True

document.save('./text/result_footers.docx')
