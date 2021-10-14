"""
СКРЫТЬ ТЕКСТ, НЕ ВЫДЕЛЕННЫЙ КУРСИВОМ
ЕСЛИ ТЕКСТ, ВЫДЕЛЕННЫЙ КУРСИВОМ, БЫЛ УЖЕ СКРЫТЫЙ, ОН ОСТАНЕТСЯ СКРЫТЫМ

не скрывает автоматически сгенерированное содержание и некоторые кривые списки
"""
from docx import Document, opc

try:
    document = Document('./text/italic.docx')
except opc.exceptions.PackageNotFoundError:
    raise Exception("Document is empty or doesn't exist")

# скрываем некурсивный текст везде кроме колонтитулов
for t in document.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.italic is None:
                        run.font.hidden = True

for p in document.paragraphs:
    for run in p.runs:

        if run.font.italic is None:
            run.font.hidden = True



def iterate_footer_header(footer_or_header):
    """
    используем методы  docx библиотеки, чтобы проверить весь текст в колонтитулах.
    работаем с таблицами и с параграфами, которые могут быть в колонтитулах.
    """
    for t in footer_or_header.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.italic is None:
                            run.font.hidden = True
    for p in footer_or_header.paragraphs:
        for run in p.runs:
            if run.font.italic is None:
                run.font.hidden = True




# итерируемся по всем секциям в документе и обрабатываем все три вида верхних и нижних колонтитулов.
# боковые колонтитулы тоже входят

for section in document.sections:

    footers_and_headers = (section.footer, section.even_page_footer, section.first_page_footer,
                           section.header, section.even_page_header, section.first_page_header)
    for footer in footers_and_headers:
        iterate_footer_header(footer)

document.save('./text/done/italic_pydocx.docx')
