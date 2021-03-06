"""
FOR ELECTRON

СКРЫТЬ ТЕКСТ, _НЕ_ ВЫДЕЛЕННЫЙ ОДНИМ ИЗ: КУРСИВ, ЦВЕТ, ЖИРНЫЙ ШРИФТ
ЕСЛИ ВЫДЕЛЕННЫЙ ТЕКСТ УЖЕ БЫЛ СКРЫТЫМ, ОН ОСТАНЕТСЯ СКРЫТЫМ

не скрывает автоматически сгенерированное содержание и некоторые кривые списки
"""
from docx import Document, opc

italic = "italic"
highlighted = "highlight_color"


def hide(path_to_file_with_name: str, file_name: str, task_list: list):
    utils_py.log_in_file(f"HIGHLIGHTED: {path_to_file_with_name} {file_name}")

    try:
        document = Document(path_to_file_with_name)
    except opc.exceptions.PackageNotFoundError:
        utils_py.log_in_file(f"HIGHLIGHTED: PackageNotFoundError {path_to_file_with_name} {file_name}")

        Document().save(f'{constants.complete_path}{file_name}')
        return

    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if italic in task_list:
                            if run.font.italic is None:
                                run.font.hidden = True
                        if highlighted in task_list:
                            if run.font.highlight_color is None:
                                run.font.hidden = True

    for p in document.paragraphs:
        for run in p.runs:
            if italic in task_list:
                if run.font.italic is None:
                    run.font.hidden = True
            if highlighted in task_list:
                if run.font.highlight_color is None:
                    run.font.hidden = True

    def iterate_footer_header(footer_or_header):
        """
        используем методы  docx библиотеки, чтобы проверить весь текст в колонтитулах.
        работаем с таблицами и с параграфами, которые могут быть в колонтитулах.
        """
        for t in footer_or_header.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if italic in task_list:
                                if run.font.italic is None:
                                    run.font.hidden = True
                            if highlighted in task_list:
                                if run.font.highlight_color is None:
                                    run.font.hidden = True

        for p in footer_or_header.paragraphs:
            for run in p.runs:
                if italic in task_list:
                    if run.font.italic is None:
                        run.font.hidden = True
                if highlighted in task_list:
                    if run.font.highlight_color is None:
                        run.font.hidden = True

    # итерируемся по всем секциям в документе и обрабатываем все три вида верхних и нижних колонтитулов.
    # боковые колонтитулы тоже входят

    for section in document.sections:

        footers_and_headers = (section.footer, section.even_page_footer, section.first_page_footer,
                               section.header, section.even_page_header, section.first_page_header)
        for footer in footers_and_headers:
            iterate_footer_header(footer)

    document.save(f'{constants.complete_path}{file_name}')
