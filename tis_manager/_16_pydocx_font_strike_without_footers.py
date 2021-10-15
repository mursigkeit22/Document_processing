"""
не скрывает автоматически сгенерированное содержание
"""

from docx import *

document = Document('./text/зачеркнутый.docx')

for t in document.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.strike:
                        run.font.hidden = True

for p in document.paragraphs:
    for run in p.runs:

        if run.font.strike:
            run.font.hidden = True

document.save('./text/done/зачеркнутый_done.docx')
