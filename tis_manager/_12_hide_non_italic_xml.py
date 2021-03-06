"""
СКРЫТЬ ТЕКСТ, НЕ ВЫДЕЛЕННЫЙ КУРСИВОМ
ЕСЛИ ТЕКСТ, ВЫДЕЛЕННЫЙ КУРСИВОМ, БЫЛ УЖЕ СКРЫТЫЙ, ОН ОСТАНЕТСЯ СКРЫТЫМ
<rPr> Structured Document Tag End Character Run Properties
The <rPr> elements under the tag's properties specify that
this structured document tag specifies that its start character
shall have formatting in the character style UserName,
 and that the end character shall have the formatting in the character style UserName
  as well as bold and italic direct formatting.

<pPr> (Text Paragraph Properties)
This element contains all paragraph level text properties for the containing paragraph.
These paragraph properties should override any and all conflicting properties
that are associated with the paragraph in question.
"""

from docx import oxml, Document, opc

try:
    document = Document('./text/cursive_only.docx')
except opc.exceptions.PackageNotFoundError:
    raise Exception("Document is empty or doesn't exist")

# текст документа
elements_list = document._element.xpath('//w:r')
# разделы документа (в которых задаются колонтитулы и прочие атрибуты, применяемые ко всему разделу)
sections_list = document._element.xpath('//w:sectPr')

WPML_URI = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
ODML_URI = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
tag_rPr = WPML_URI + 'rPr'
tag_r = WPML_URI + 'r'
tag_highlight = WPML_URI + 'highlight'
tag_t = WPML_URI + 't'  # текст
tag_vanish = WPML_URI + 'vanish'
tag_italic = WPML_URI + 'i'  # <w:i/>


def hide_not_highlighted_text(elements):
    """
    не используем стандартные методы библиотеки docx,
    так как с их помощью нельзя скрыть автоматически сгенерированное содержание.

    """
    for element in elements:
        print(element.xml)
        print("++++++++++++++++")
        try:
            rpr_list = element.findall(tag_rPr)  # теги, относящиеся к параграфу
            if len(rpr_list) == 0:
                tag_text = element.find(tag_t)
                rpr = oxml.shared.OxmlElement('w:rPr')  # создаем тег rPr
                rpr.append(oxml.shared.OxmlElement('w:vanish'))  # добавляем в него тег скрытого текста
                tag_text.addprevious(rpr)  # добавляем тег форматирования параграфа выше тега с текстом

            for rPr in rpr_list:  # если уже есть тег rPr, т.е. у параграфа уже есть форматирование, то смотрим на это форматирование:
                ital = rPr.findall(tag_italic)  # проверяем, есть ли тег, отвечающий за подсветку
                if len(ital) == 0:
                    already_hidden = rPr.findall(tag_vanish)
                    if len(already_hidden) == 0:
                        rPr.append(oxml.shared.OxmlElement('w:vanish'))  # если нет, добавляем в него тег скрытого текста

        except AttributeError:  # if there is no text
            pass


def iterate_footer_header(footer_or_header):
    """
    используем методы  docx библиотеки, чтобы проверить весь текст в колонтитулах.
    работаем с таблицами и с параграфами, которые могут быть в колонтитулах.
    """
    for t in footer_or_header.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.italic is None:
                            run.font.hidden = True
    for p in footer_or_header.paragraphs:
        for run in p.runs:
            if run.font.italic is None:
                run.font.hidden = True


# скрываем невыделенный текст везде кроме колонтитулов
hide_not_highlighted_text(elements_list)

# итерируемся по всем секциям в документе и обрабатываем все три вида верхних и нижних колонтитулов.
# боковые колонтитулы тоже входят

for section in document.sections:

    footers_and_headers = (section.footer, section.even_page_footer, section.first_page_footer,
                           section.header, section.even_page_header, section.first_page_header)
    for footer in footers_and_headers:
        iterate_footer_header(footer)

document.save('./text/done/cursive_only_done.docx')
